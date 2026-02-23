"""
NativeParserSkill - 原生文档解析器

负责解析以下类型的文档：
1. 纯文本文件 (.txt)
2. Markdown 文件 (.md)
3. Word 文档 (.docx) - 保留标题层级
4. 原生 PDF (.pdf) - 保持段落逻辑
"""

import asyncio
from pathlib import Path
from typing import Dict, Any, List
from .base_skill import BaseSkill


class NativeParserSkill(BaseSkill):
    """
    原生文档解析 Skill

    支持 TXT、Markdown、Word、原生 PDF 的解析
    """

    def __init__(self, preserve_formatting: bool = True):
        """
        初始化 NativeParserSkill

        Args:
            preserve_formatting: 是否保留格式（标题层级、段落等）
        """
        super().__init__(name="NativeParser")
        self.preserve_formatting = preserve_formatting

    async def execute(self, input_data: Any, **kwargs) -> Dict[str, Any]:
        """
        执行文档解析

        Args:
            input_data: 文件路径（str 或 Path）或路由结果字典
            **kwargs: 额外参数

        Returns:
            {
                'file_path': str,
                'file_type': str,
                'content': str,           # 解析后的文本内容
                'metadata': dict,         # 文档元数据
                'structure': list         # 文档结构（如标题层级）
            }

        Raises:
            ValueError: 不支持的文件类型
            FileNotFoundError: 文件不存在
        """
        # 处理输入数据
        if isinstance(input_data, dict):
            file_path = Path(input_data['file_path'])
            file_type = input_data.get('file_type')
        else:
            file_path = Path(input_data)
            file_type = None

        # 检查文件
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        # 根据文件类型选择解析方法
        file_ext = file_path.suffix.lower()

        if file_ext in ['.txt', '.md']:
            result = await self._parse_text(file_path, file_ext)
        elif file_ext == '.docx':
            result = await self._parse_word(file_path)
        elif file_ext == '.pdf':
            result = await self._parse_pdf(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {file_ext}")

        return result

    async def _parse_text(self, file_path: Path, file_ext: str) -> Dict[str, Any]:
        """
        解析纯文本和 Markdown 文件

        Args:
            file_path: 文件路径
            file_ext: 文件扩展名

        Returns:
            解析结果字典
        """
        self.logger.info(f"📝 解析文本文件: {file_path.name}")

        # 读取文件内容
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # 提取结构（Markdown 标题）
        structure = []
        if file_ext == '.md' and self.preserve_formatting:
            structure = self._extract_markdown_structure(content)

        return {
            'file_path': str(file_path),
            'file_type': 'markdown' if file_ext == '.md' else 'text',
            'content': content,
            'metadata': {
                'file_name': file_path.name,
                'file_size': file_path.stat().st_size,
                'line_count': content.count('\n') + 1,
                'char_count': len(content)
            },
            'structure': structure
        }

    def _extract_markdown_structure(self, content: str) -> List[Dict[str, Any]]:
        """
        提取 Markdown 文档的标题结构

        Args:
            content: Markdown 内容

        Returns:
            标题结构列表
        """
        structure = []
        lines = content.split('\n')

        for i, line in enumerate(lines):
            line = line.strip()
            if line.startswith('#'):
                # 计算标题级别
                level = 0
                for char in line:
                    if char == '#':
                        level += 1
                    else:
                        break

                # 提取标题文本
                title = line[level:].strip()
                structure.append({
                    'level': level,
                    'title': title,
                    'line_number': i + 1
                })

        return structure

    async def _parse_word(self, file_path: Path) -> Dict[str, Any]:
        """
        解析 Word 文档（保留标题层级）

        Args:
            file_path: Word 文件路径

        Returns:
            解析结果字典
        """
        self.logger.info(f"📄 解析 Word 文档: {file_path.name}")

        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "需要安装 python-docx 库来解析 Word 文档\n"
                "安装方法: pip install python-docx"
            )

        # 打开文档
        doc = Document(file_path)

        # 提取内容和结构
        content_parts = []
        structure = []

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            # 检查是否为标题
            if para.style.name.startswith('Heading'):
                # 提取标题级别
                try:
                    level = int(para.style.name.split()[-1])
                except:
                    level = 1

                structure.append({
                    'level': level,
                    'title': text,
                    'paragraph_index': i
                })

                # 保留格式：添加 Markdown 风格的标题标记
                if self.preserve_formatting:
                    content_parts.append(f"\n{'#' * level} {text}\n")
                else:
                    content_parts.append(text)
            else:
                content_parts.append(text)

        content = '\n'.join(content_parts)

        return {
            'file_path': str(file_path),
            'file_type': 'word',
            'content': content,
            'metadata': {
                'file_name': file_path.name,
                'file_size': file_path.stat().st_size,
                'paragraph_count': len(doc.paragraphs),
                'char_count': len(content)
            },
            'structure': structure
        }

    async def _parse_pdf(self, file_path: Path) -> Dict[str, Any]:
        """
        解析原生 PDF（保持段落逻辑）

        Args:
            file_path: PDF 文件路径

        Returns:
            解析结果字典
        """
        self.logger.info(f"📕 解析 PDF 文档: {file_path.name}")

        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError(
                "需要安装 PyMuPDF 库来解析 PDF 文档\n"
                "安装方法: pip install PyMuPDF"
            )

        # 打开 PDF
        doc = fitz.open(file_path)

        # 提取文本
        content_parts = []
        structure = []

        for page_num in range(len(doc)):
            page = doc[page_num]

            # 提取文本（保持布局）
            if self.preserve_formatting:
                text = page.get_text("text")
            else:
                text = page.get_text()

            content_parts.append(text)

            # 尝试提取标题（基于字体大小）
            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if "lines" in block:
                    for line in block["lines"]:
                        for span in line["spans"]:
                            # 如果字体较大，可能是标题
                            if span["size"] > 14:  # 字体大小阈值
                                structure.append({
                                    'page': page_num + 1,
                                    'text': span["text"],
                                    'font_size': span["size"]
                                })

        doc.close()

        content = '\n\n'.join(content_parts)

        return {
            'file_path': str(file_path),
            'file_type': 'pdf',
            'content': content,
            'metadata': {
                'file_name': file_path.name,
                'file_size': file_path.stat().st_size,
                'page_count': len(doc),
                'char_count': len(content)
            },
            'structure': structure
        }


# 使用示例
if __name__ == "__main__":
    import asyncio

    async def test_parser():
        """测试文档解析功能"""
        parser = NativeParserSkill(preserve_formatting=True)

        # 测试文件
        test_files = [
            "data/example.txt",
            "data/example.md",
            "data/example.docx",
            "data/example.pdf",
        ]

        for file_path in test_files:
            print(f"\n{'='*60}")
            print(f"测试文件: {file_path}")
            print('='*60)

            result = await parser.run(file_path)

            if result['success']:
                data = result['data']
                print(f"✅ 解析成功:")
                print(f"   文件类型: {data['file_type']}")
                print(f"   字符数: {data['metadata']['char_count']}")
                print(f"   结构元素: {len(data['structure'])}")
                print(f"   内容预览: {data['content'][:200]}...")
            else:
                print(f"❌ 解析失败: {result['error']}")

        # 查看统计信息
        print(f"\n{'='*60}")
        print("统计信息:")
        print('='*60)
        stats = parser.get_stats()
        for key, value in stats.items():
            print(f"   {key}: {value}")

    # 运行测试
    asyncio.run(test_parser())
